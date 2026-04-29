import os
import json
import glob
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define Error Taxonomy with Descriptions and Examples
ERROR_TAXONOMY = {
    "pronoun_failure": {
        "desc": "Model failed to resolve a pronoun (e.g., 'it', 'them') to the correct object mentioned earlier.",
        "example": "Instruction: 'Pick up the apple. Heat it.' -> Agent failed to associate 'it' with the apple."
    },
    "attributive_misid": {
        "desc": "Wrong object picked based on descriptions like color, size, or state.",
        "example": "Instruction: 'Pick up the red apple.' -> Agent picks up a 'green apple' instead."
    },
    "noise_distraction": {
        "desc": "Ambiguous names or irrelevant objects in the scene caused the model to target the wrong thing.",
        "example": "Instruction: 'Find the book.' -> Agent targets 'book_shelf' or an irrelevant 'newspaper'."
    },
    "context_loss": {
        "desc": "Failure to maintain task context across multiple steps, leading to hallucinations.",
        "example": "Agent starts a task but forgets the target halfway through and targets a random object."
    },
    "action_ordering": {
        "desc": "Required actions were present but performed in the wrong logical sequence.",
        "example": "Agent tries to 'pick_up' the apple before using 'find' to locate it."
    },
    "action_omission": {
        "desc": "Missing critical steps required by the ground truth plan.",
        "example": "Instruction: 'Heat and serve the potato.' -> Agent picks and places but forgets to 'heat'."
    },
    "hallucinated_action": {
        "desc": "Agent attempted actions that are not valid skills or involve non-existent objects.",
        "example": "Agent generates a 'cook' command when only 'heat' is available."
    },
    "wrong_destination": {
        "desc": "Object was correctly manipulated but placed at the wrong location.",
        "example": "Instruction: 'Put it in the fridge.' -> Agent puts it on the 'counter'."
    }
}

def json_to_word(json_path):
    model_dir = os.path.dirname(json_path)
    model_name = os.path.basename(model_dir)
    docx_path = os.path.join(model_dir, f"{model_name}_analysis_report.docx")
    
    print(f"Generating Word report for {model_name}...")
    
    with open(json_path, 'r') as f:
        analysis_data = json.load(f)
        
    # Load original results to get instruction/plans (since they aren't in the judge json)
    results_files = glob.glob(os.path.join(model_dir, "results_*.json"))
    if not results_files:
        print(f"Warning: No results file found for {model_name} to pull context data.")
        return
        
    with open(results_files[0], 'r') as f:
        original_data = json.load(f)
        
    # Build a lookup for examples
    example_lookup = {}
    for condition, results_list in original_data.get("detailed_results", {}).items():
        for res in results_list:
            example_lookup[res["example_id"]] = res

    doc = Document()
    
    # Title
    title = doc.add_heading(f'REI-Bench Failure Analysis: {model_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Taxonomy Section (Glossary)
    doc.add_heading('1. Error Taxonomy & Examples', level=1)
    doc.add_paragraph("This section defines the fine-grained error categories used by the GPT-5 judge to classify failures.")
    
    for cat, info in ERROR_TAXONOMY.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{cat.replace('_', ' ').title()}: ").bold = True
        p.add_run(info["desc"])
        
        example_p = doc.add_paragraph()
        example_p.paragraph_format.left_indent = Inches(0.5)
        example_p.add_run("Example: ").italic = True
        example_p.add_run(info["example"]).italic = True
    
    doc.add_page_break()

    # Summary Section
    doc.add_heading('2. Error Category Distribution', level=1)
    doc.add_paragraph(f"Total Failures Analyzed: {analysis_data.get('total_failures', 'N/A')}")
    
    # Table for percentages
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Error Category'
    hdr_cells[1].text = 'Percentage'
    
    percentages = analysis_data.get("error_category_percentages", {})
    for cat, pct in sorted(percentages.items(), key=lambda x: -x[1]):
        row_cells = table.add_row().cells
        row_cells[0].text = cat.replace('_', ' ').title()
        row_cells[1].text = f"{pct}%"
        
    doc.add_page_break()
    
    # Detailed Section
    doc.add_heading('3. Detailed Failure Case Analysis', level=1)
    
    for eval_item in analysis_data.get("detailed_evaluations", []):
        eid = eval_item["example_id"]
        ctx = example_lookup.get(eid, {})
        
        doc.add_heading(f"Example: {eid}", level=2)
        
        # Metadata
        p = doc.add_paragraph()
        p.add_run("Task Type: ").bold = True
        p.add_run(f"{eval_item['task_type']}")
        
        p = doc.add_paragraph()
        p.add_run("LLM Judge Category: ").bold = True
        p.add_run(f"{eval_item['error_category'].replace('_', ' ').title()}")
        
        # Instruction
        doc.add_heading("Instruction:", level=3)
        doc.add_paragraph(ctx.get("instruction", "N/A"))
        
        # Reasoning
        doc.add_heading("GPT-5 Judge Reasoning:", level=3)
        reason_p = doc.add_paragraph(eval_item.get("reason", "N/A"))
        reason_p.style.font.italic = True
        
        # Plans Comparison Table
        plan_table = doc.add_table(rows=1, cols=2)
        plan_table.style = 'Table Grid'
        plan_hdr = plan_table.rows[0].cells
        plan_hdr[0].text = 'Ground Truth Plan'
        plan_hdr[1].text = 'Generated Plan (Failure)'
        
        row = plan_table.add_row().cells
        row[0].text = "\n".join([f"- {a}" for a in ctx.get("ground_truth_plan", [])])
        row[1].text = "\n".join([f"- {a}" for a in ctx.get("generated_plan", [])])
        
        doc.add_paragraph("\n") # Spacer

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

def main():
    judge_jsons = glob.glob("phase1/results/*/*_llmjudge.json")
    if not judge_jsons:
        print("No *_llmjudge.json files found.")
        return
        
    for j_path in judge_jsons:
        json_to_word(j_path)

if __name__ == "__main__":
    main()
